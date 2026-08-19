"""
DOCX report generator for test evidence documentation.
"""
import os
from datetime import datetime
from typing import Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

from session_model import TestSession


class ReportGenerator:
    """Generates structured Word documents from test sessions."""
    
    def __init__(self):
        """Initialize report generator."""
        pass
    
    def generate(self, session: TestSession, output_dir: str = None) -> str:
        """
        Generate DOCX evidence report from test session.
        
        Args:
            session: Completed TestSession with steps
            output_dir: Directory to save the report (defaults to user's Downloads folder)
            
        Returns:
            Path to generated DOCX file
            
        Raises:
            Exception: If report generation fails
        """
        try:
            # Use Downloads folder if no output_dir specified
            if output_dir is None or output_dir == "./output":
                # Get user's Downloads folder
                if os.name == 'nt':  # Windows
                    import ctypes
                    from ctypes import wintypes
                    
                    # Get Downloads folder path on Windows
                    CSIDL_PERSONAL = 5  # My Documents
                    CSIDL_DOWNLOADS = 0x0028  # Downloads folder
                    
                    # Try to get Downloads folder
                    try:
                        buffer = ctypes.create_unicode_buffer(wintypes.MAX_PATH)
                        ctypes.windll.shell32.SHGetFolderPathW(0, CSIDL_DOWNLOADS, 0, 0, buffer)
                        output_dir = buffer.value
                    except:
                        # Fallback to Documents folder if Downloads not available
                        try:
                            ctypes.windll.shell32.SHGetFolderPathW(0, CSIDL_PERSONAL, 0, 0, buffer)
                            output_dir = buffer.value
                        except:
                            # Last resort fallback
                            output_dir = os.path.expanduser("~/Downloads")
                else:
                    # macOS and Linux
                    output_dir = os.path.expanduser("~/Downloads")
            
            # Create output directory if it doesn't exist
            os.makedirs(output_dir, exist_ok=True)
            
            # Generate filename
            date_str = datetime.now().strftime("%Y%m%d")
            filename = f"Evidence_{session.tc_id}_{date_str}_{session.tester_name}.docx"
            output_path = os.path.join(output_dir, filename)
            
            # Create document
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            # Generate report sections
            self._add_cover_page(doc, session)
            doc.add_page_break()
            
            self._add_summary_section(doc, session)
            doc.add_page_break()
            
            self._add_step_evidence_section(doc, session)
            
            self._add_signoff_block(doc, session)
            
            # Save document
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"Report generation failed: {str(e)}")
    
    def _add_cover_page(self, doc: Document, session: TestSession) -> None:
        """Add cover page with test metadata."""
        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("TestTrace Recorder\nTest Evidence Report")
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(37, 99, 235)  # #2563EB
        
        doc.add_paragraph()  # Spacing
        
        # Metadata table
        table = doc.add_table(rows=9, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Light Grid Accent 1'
        
        # Define metadata rows
        metadata = [
            ("Test Case ID", session.tc_id),
            ("Test Case Name", session.tc_name),
            ("Module / Feature", session.module),
            ("Environment", session.environment),
            ("Tester Name", session.tester_name),
            ("Execution Date", session.start_time.strftime("%d-%b-%Y")),
            ("Execution Time", session.start_time.strftime("%I:%M:%S %p")),
            ("Total Duration", session.get_duration()),
            ("Overall Status", session.get_overall_status())
        ]
        
        for i, (label, value) in enumerate(metadata):
            # Label cell
            label_cell = table.rows[i].cells[0]
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.font.bold = True
            label_run.font.size = Pt(11)
            
            # Value cell
            value_cell = table.rows[i].cells[1]
            value_para = value_cell.paragraphs[0]
            value_run = value_para.add_run(str(value))
            value_run.font.size = Pt(11)
            
            # Color-code status
            if label == "Overall Status":
                if value == "PASS":
                    value_run.font.color.rgb = RGBColor(22, 163, 74)  # Green
                    value_run.font.bold = True
                elif value == "FAIL":
                    value_run.font.color.rgb = RGBColor(220, 38, 38)  # Red
                    value_run.font.bold = True
                elif value == "BLOCKED":
                    value_run.font.color.rgb = RGBColor(245, 158, 11)  # Amber
                    value_run.font.bold = True
    
    def _add_summary_section(self, doc: Document, session: TestSession) -> None:
        """Add test execution summary."""
        # Section heading
        heading = doc.add_heading("Test Execution Summary", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Result counts
        counts = session.get_result_counts()
        
        summary_text = (
            f"Total Steps Executed: {len(session.steps)}\n"
            f"Passed: {counts['Pass']}\n"
            f"Failed: {counts['Fail']}\n"
            f"Blocked: {counts['Blocked']}\n"
            f"Untested: {counts['Untested']}\n\n"
            f"Overall Test Status: {session.get_overall_status()}"
        )
        
        para = doc.add_paragraph(summary_text)
        para.style = 'Body Text'
    
    def _add_step_evidence_section(self, doc: Document, session: TestSession) -> None:
        """Add detailed step-by-step evidence."""
        # Section heading
        heading = doc.add_heading("Step-by-Step Evidence", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        doc.add_paragraph()
        
        # Add each step
        for i, step in enumerate(session.steps):
            self._add_step_block(doc, step)
            
            # Add page break after every 2 steps (except last)
            if (i + 1) % 2 == 0 and i < len(session.steps) - 1:
                doc.add_page_break()
    
    def _add_step_block(self, doc: Document, step) -> None:
        """Add a single step evidence block."""
        # Header table (step info)
        header_table = doc.add_table(rows=1, cols=3)
        header_table.style = 'Light Grid Accent 1'
        
        # Apply blue background to header row
        header_cells = header_table.rows[0].cells
        for cell in header_cells:
            cell._element.get_or_add_tcPr().append(
                self._create_cell_shading(RGBColor(37, 99, 235))
            )
        
        # Step number
        step_cell = header_cells[0]
        step_para = step_cell.paragraphs[0]
        step_run = step_para.add_run(f"Step {step.step_number}")
        step_run.font.bold = True
        step_run.font.size = Pt(12)
        step_run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Timestamp
        time_cell = header_cells[1]
        time_para = time_cell.paragraphs[0]
        time_run = time_para.add_run(step.timestamp)
        time_run.font.size = Pt(10)
        time_run.font.color.rgb = RGBColor(255, 255, 255)
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Active window
        window_cell = header_cells[2]
        window_para = window_cell.paragraphs[0]
        window_run = window_para.add_run(step.active_window[:50])  # Truncate long titles
        window_run.font.size = Pt(9)
        window_run.font.italic = True
        window_run.font.color.rgb = RGBColor(255, 255, 255)
        window_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Description
        desc_para = doc.add_paragraph()
        desc_run = desc_para.add_run(f"Action: {step.description}")
        desc_run.font.size = Pt(11)
        
        doc.add_paragraph()  # Spacing
        
        # Screenshot
        screenshot_path = step.annotated_path or step.screenshot_path
        if os.path.exists(screenshot_path):
            try:
                # Add image with max width
                doc.add_picture(screenshot_path, width=Inches(6.0))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                doc.add_paragraph(f"[Screenshot not available: {e}]")
        else:
            doc.add_paragraph("[Screenshot file not found]")
        
        doc.add_paragraph()  # Spacing
        
        # Result badge
        result_para = doc.add_paragraph()
        result_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        result_run = result_para.add_run(f"  {step.result.upper()}  ")
        result_run.font.size = Pt(12)
        result_run.font.bold = True
        
        # Color-code result using highlight color (since shading isn't available on ParagraphFormat)
        if step.result == "Pass":
            result_run.font.color.rgb = RGBColor(255, 255, 255)
            # Apply shading to the run
            from docx.oxml import OxmlElement
            shd = OxmlElement('w:shd')
            shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', '16A34A')
            result_run._element.get_or_add_rPr().append(shd)
        elif step.result == "Fail":
            result_run.font.color.rgb = RGBColor(255, 255, 255)
            # Red background
            from docx.oxml import OxmlElement
            shd = OxmlElement('w:shd')
            shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', 'DC2626')
            result_run._element.get_or_add_rPr().append(shd)
        elif step.result == "Blocked":
            result_run.font.color.rgb = RGBColor(0, 0, 0)
            # Amber background
            from docx.oxml import OxmlElement
            shd = OxmlElement('w:shd')
            shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', 'F59E0B')
            result_run._element.get_or_add_rPr().append(shd)
        
        # Separator line
        doc.add_paragraph("_" * 80)
        doc.add_paragraph()
    
    def _add_signoff_block(self, doc: Document, session: TestSession) -> None:
        """Add tester sign-off section."""
        doc.add_page_break()
        
        heading = doc.add_heading("Sign-Off", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Sign-off table
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Tester name
        table.rows[0].cells[0].text = "Tester Name:"
        table.rows[0].cells[1].text = session.tester_name
        
        # Date
        table.rows[1].cells[0].text = "Date:"
        table.rows[1].cells[1].text = datetime.now().strftime("%d-%b-%Y")
        
        # Signature
        table.rows[2].cells[0].text = "Signature:"
        table.rows[2].cells[1].text = "_" * 40
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Reviewer section
        doc.add_paragraph("Reviewed By:")
        doc.add_paragraph()
        
        reviewer_table = doc.add_table(rows=3, cols=2)
        reviewer_table.style = 'Light Grid Accent 1'
        
        reviewer_table.rows[0].cells[0].text = "Reviewer Name:"
        reviewer_table.rows[0].cells[1].text = ""
        
        reviewer_table.rows[1].cells[0].text = "Review Date:"
        reviewer_table.rows[1].cells[1].text = ""
        
        reviewer_table.rows[2].cells[0].text = "Signature:"
        reviewer_table.rows[2].cells[1].text = "_" * 40
    
    @staticmethod
    def _create_cell_shading(color: RGBColor):
        """Create XML element for cell background color."""
        from docx.oxml import OxmlElement
        
        shading_elm = OxmlElement('w:shd')
        # RGBColor is actually a tuple (r, g, b), not an object with attributes
        r, g, b = color if isinstance(color, tuple) else (color[0], color[1], color[2])
        shading_elm.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill',
                       f"{r:02X}{g:02X}{b:02X}")
        return shading_elm
