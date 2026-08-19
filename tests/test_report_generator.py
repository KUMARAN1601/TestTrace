"""
Unit tests for report_generator.py - ReportGenerator class.
"""
import pytest
import os
import tempfile
from datetime import datetime
from docx import Document
from report_generator import ReportGenerator
from session_model import TestSession, TestStep


class TestReportGenerator:
    """Test cases for ReportGenerator class."""
    
    def test_generator_initialization(self):
        """Test ReportGenerator instantiation."""
        generator = ReportGenerator()
        
        assert generator is not None
    
    def test_generate_report_creates_file(self):
        """Test that report generation creates a DOCX file."""
        generator = ReportGenerator()
        
        # Create mock session
        session = TestSession(
            tc_id="TC_RPT_001",
            tc_name="Test Report Generation",
            module="Testing",
            environment="SIT",
            tester_name="Test User"
        )
        
        # Add a test step (without screenshot for now)
        step = TestStep(
            step_number=1,
            timestamp="2026-08-19 10:30:00",
            screenshot_path="",
            description="Test step for report",
            result="Pass"
        )
        session.add_step(step)
        
        # Generate report in temp directory
        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = generator.generate(session, temp_dir)
            
            assert os.path.exists(output_path)
            assert output_path.endswith(".docx")
            assert "Evidence_TC_RPT_001" in output_path
    
    def test_generated_report_structure(self):
        """Test that generated report has correct structure."""
        generator = ReportGenerator()
        
        # Create mock session with multiple steps
        session = TestSession(
            tc_id="TC_RPT_002",
            tc_name="Test Report Structure",
            module="Core",
            environment="UAT",
            tester_name="QA Tester"
        )
        
        # Add multiple steps with different results
        steps_data = [
            ("Step 1 description", "Pass"),
            ("Step 2 description", "Pass"),
            ("Step 3 description", "Fail"),
            ("Step 4 description", "Blocked"),
        ]
        
        for i, (desc, result) in enumerate(steps_data, start=1):
            step = TestStep(
                step_number=i,
                timestamp=f"2026-08-19 10:3{i}:00",
                description=desc,
                result=result,
                active_window="Test Application"
            )
            session.add_step(step)
        
        # Generate report
        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = generator.generate(session, temp_dir)
            
            # Open and verify document structure
            doc = Document(output_path)
            
            # Check that document has content
            assert len(doc.paragraphs) > 0
            
            # Check for key sections (looking for specific text in paragraphs and tables)
            doc_text = "\n".join([p.text for p in doc.paragraphs])
            
            # Also check table content (cover page has metadata in a table)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        doc_text += "\n" + cell.text
            
            assert "TestTrace Recorder" in doc_text
            assert "TC_RPT_002" in doc_text
            assert "Test Report Structure" in doc_text
            assert "QA Tester" in doc_text
    
    def test_report_with_empty_session(self):
        """Test report generation with no steps."""
        generator = ReportGenerator()
        
        session = TestSession(
            tc_id="TC_RPT_003",
            tc_name="Empty Session",
            module="Test",
            environment="SIT",
            tester_name="Tester"
        )
        
        # No steps added
        
        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = generator.generate(session, temp_dir)
            
            assert os.path.exists(output_path)
            
            # Document should still be created with cover page
            doc = Document(output_path)
            assert len(doc.paragraphs) > 0
    
    def test_report_filename_format(self):
        """Test that report filename follows correct format."""
        generator = ReportGenerator()
        
        session = TestSession(
            tc_id="TC_VISA_AUTH_001",
            tc_name="VISA Authorization Test",
            module="Authorization",
            environment="PROD",
            tester_name="Kumaran"
        )
        
        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = generator.generate(session, temp_dir)
            
            filename = os.path.basename(output_path)
            
            # Check filename format: Evidence_{tc_id}_{date}_{tester_name}.docx
            assert filename.startswith("Evidence_TC_VISA_AUTH_001_")
            assert filename.endswith("_Kumaran.docx")
            assert "2026" in filename  # Should contain current year
    
    def test_report_with_all_result_types(self):
        """Test report generation with all result types."""
        generator = ReportGenerator()
        
        session = TestSession(
            tc_id="TC_RPT_004",
            tc_name="All Results Test",
            module="Testing",
            environment="SIT",
            tester_name="Tester"
        )
        
        results = ["Pass", "Fail", "Blocked", "Untested"]
        for i, result in enumerate(results, start=1):
            step = TestStep(
                step_number=i,
                description=f"Step with {result} result",
                result=result
            )
            session.add_step(step)
        
        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = generator.generate(session, temp_dir)
            
            assert os.path.exists(output_path)
            
            # Verify document contains all result types
            doc = Document(output_path)
            doc_text = "\n".join([p.text for p in doc.paragraphs])
            
            assert "PASS" in doc_text
            assert "FAIL" in doc_text
            assert "BLOCKED" in doc_text
    
    def test_overall_status_calculation(self):
        """Test that overall status is correctly calculated in report."""
        generator = ReportGenerator()
        
        # Test PASS scenario (all pass)
        session_pass = TestSession(
            tc_id="TC_RPT_005",
            tc_name="Pass Test",
            module="Test",
            environment="SIT",
            tester_name="Tester"
        )
        session_pass.add_step(TestStep(step_number=1, result="Pass"))
        session_pass.add_step(TestStep(step_number=2, result="Pass"))
        
        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = generator.generate(session_pass, temp_dir)
            doc = Document(output_path)
            doc_text = "\n".join([p.text for p in doc.paragraphs])
            
            # Should show PASS in overall status
            assert "PASS" in doc_text
        
        # Test FAIL scenario (any fail)
        session_fail = TestSession(
            tc_id="TC_RPT_006",
            tc_name="Fail Test",
            module="Test",
            environment="SIT",
            tester_name="Tester"
        )
        session_fail.add_step(TestStep(step_number=1, result="Pass"))
        session_fail.add_step(TestStep(step_number=2, result="Fail"))
        
        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = generator.generate(session_fail, temp_dir)
            doc = Document(output_path)
            doc_text = "\n".join([p.text for p in doc.paragraphs])
            
            # Should show FAIL in overall status
            assert "FAIL" in doc_text
    
    def test_result_counts_in_summary(self):
        """Test that result counts are included in summary."""
        generator = ReportGenerator()
        
        session = TestSession(
            tc_id="TC_RPT_007",
            tc_name="Count Test",
            module="Test",
            environment="SIT",
            tester_name="Tester"
        )
        
        # Add 3 pass, 1 fail, 1 blocked
        session.add_step(TestStep(step_number=1, result="Pass"))
        session.add_step(TestStep(step_number=2, result="Pass"))
        session.add_step(TestStep(step_number=3, result="Pass"))
        session.add_step(TestStep(step_number=4, result="Fail"))
        session.add_step(TestStep(step_number=5, result="Blocked"))
        
        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = generator.generate(session, temp_dir)
            doc = Document(output_path)
            doc_text = "\n".join([p.text for p in doc.paragraphs])
            
            # Check counts are present
            assert "Total Steps Executed: 5" in doc_text
            assert "Passed: 3" in doc_text
            assert "Failed: 1" in doc_text
            assert "Blocked: 1" in doc_text
    
    def test_output_directory_created(self):
        """Test that output directory is created if it doesn't exist."""
        generator = ReportGenerator()
        
        session = TestSession(
            tc_id="TC_RPT_008",
            tc_name="Dir Test",
            module="Test",
            environment="SIT",
            tester_name="Tester"
        )
        
        with tempfile.TemporaryDirectory() as temp_dir:
            # Use a subdirectory that doesn't exist
            output_dir = os.path.join(temp_dir, "new_folder", "reports")
            
            output_path = generator.generate(session, output_dir)
            
            # Directory should be created
            assert os.path.exists(output_dir)
            assert os.path.exists(output_path)
    
    def test_generate_with_missing_screenshot(self):
        """Test report generation when screenshot file is missing."""
        generator = ReportGenerator()
        
        session = TestSession(
            tc_id="TC_RPT_009",
            tc_name="Missing Screenshot",
            module="Test",
            environment="SIT",
            tester_name="Tester"
        )
        
        # Add step with non-existent screenshot path
        step = TestStep(
            step_number=1,
            screenshot_path="/nonexistent/path/screenshot.png",
            annotated_path="/nonexistent/path/annotated.png",
            description="Step with missing screenshot",
            result="Pass"
        )
        session.add_step(step)
        
        with tempfile.TemporaryDirectory() as temp_dir:
            # Should not raise exception
            output_path = generator.generate(session, temp_dir)
            
            assert os.path.exists(output_path)
            
            # Document should contain placeholder text
            doc = Document(output_path)
            doc_text = "\n".join([p.text for p in doc.paragraphs])
            
            # Should have some indication of missing file
            assert "Step with missing screenshot" in doc_text
